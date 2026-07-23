"""Fills the DPIA DOCX template from saved structured data."""

import os
from docx import Document

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "MCP Document.docx")


def _set(cell, text: str):
    text = str(text or "").strip()
    if not text:
        return
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _set_ja_nej(cell_r1, cell_r2, value: str):
    lines = value.strip().split("\n", 1)
    _set(cell_r1, lines[0].strip())
    _set(cell_r2, lines[1].strip() if len(lines) > 1 else "")


def fill_from_data(saved: dict) -> str:
    doc = Document(TEMPLATE_PATH)
    t = doc.tables

    s1 = saved.get("steg1", {})
    s2 = saved.get("steg2", {})
    s3 = saved.get("steg3", {})
    criteria = saved.get("criteria", {})

    # ── STEG 1: Bedöm behovet (T0–T10) ──────────────────────────────────────

    _set(t[0].cell(1, 0), s1.get("overview", ""))
    _set(t[1].cell(1, 0), s1.get("faktorer", ""))
    _set(t[2].cell(1, 0), s1.get("liknar_35_3", ""))

    # T3: criteria — col 2 = Ja/Nej, col 1 = motivering
    for i in range(1, 10):
        val = criteria.get(str(i), "")
        if isinstance(val, dict):
            _set(t[3].cell(i + 1, 2), val.get("svar", ""))
            _set(t[3].cell(i + 1, 1), val.get("motivering", ""))
        else:
            _set(t[3].cell(i + 1, 2), str(val))

    _set(t[4].cell(1, 0), s1.get("tva_kriterier", ""))
    _set_ja_nej(t[5].cell(1, 0), t[5].cell(2, 0), s1.get("undantag_trots_ja", ""))
    _set_ja_nej(t[6].cell(1, 0), t[6].cell(2, 0), s1.get("genomforas_trots_nej", ""))
    _set_ja_nej(t[7].cell(1, 0), t[7].cell(2, 0), s1.get("lik_tidigare", ""))
    _set_ja_nej(t[8].cell(1, 0), t[8].cell(2, 0), s1.get("allman_35_10", ""))
    _set(t[9].cell(1, 0), s1.get("dso_svar", ""))
    _set(t[10].cell(1, 0), s1.get("behov", ""))

    # ── STEG 2: Systematisk beskrivning (T11–T21) ─────────────────────────────

    _set(t[11].cell(1, 0), s2.get("objekt", ""))

    # T12: categories of personal data (up to 6)
    for i, kat in enumerate(s2.get("kategorier", [])[:6]):
        _set(t[12].cell(i + 2, 1), kat.get("kategori", ""))
        _set(t[12].cell(i + 2, 2), kat.get("sarskild", ""))

    # T13: categories of data subjects / registrerade (up to 6)
    for i, reg in enumerate(s2.get("registrerade", [])[:6]):
        _set(t[13].cell(i + 2, 1), reg.get("kategori", ""))
        _set(t[13].cell(i + 2, 2), reg.get("sarskild", ""))

    _set(t[14].cell(1, 0), s2.get("individer", ""))
    _set(t[15].cell(1, 0), s2.get("lander", ""))
    _set(t[16].cell(1, 0), s2.get("sammanhang", ""))
    _set(t[17].cell(1, 0), s2.get("anledning", ""))
    _set(t[18].cell(1, 0), s2.get("resurser", ""))
    _set(t[19].cell(1, 0), s2.get("beskrivning", ""))
    _set(t[20].cell(1, 0), s2.get("ansvariga", ""))

    # T21: recipients — up to 6 rows
    for i, m in enumerate(s2.get("mottagare", [])[:6]):
        _set(t[21].cell(i + 2, 1), m.get("biträde", ""))
        _set(t[21].cell(i + 2, 2), m.get("kategorier", ""))
        _set(t[21].cell(i + 2, 3), m.get("ändamål", ""))

    # ── STEG 3: Rättslig analys (T22–T34) ────────────────────────────────────

    _set(t[22].cell(1, 0), s3.get("regelverk", ""))
    _set(t[23].cell(1, 0), s3.get("andamal", ""))

    # T24: legal basis — up to 10 rows
    for i, rg in enumerate(s3.get("rattslig_grund", [])[:10]):
        _set(t[24].cell(i + 2, 1), rg.get("beskrivning", ""))
        _set(t[24].cell(i + 2, 2), rg.get("personuppgifter", ""))
        _set(t[24].cell(i + 2, 3), rg.get("grund", ""))
        _set(t[24].cell(i + 2, 4), rg.get("kommentar", ""))

    # T25: special category exemptions — up to 10 rows
    for i, sg in enumerate(s3.get("sarskild_grund", [])[:10]):
        _set(t[25].cell(i + 2, 1), sg.get("beskrivning", ""))
        _set(t[25].cell(i + 2, 2), sg.get("personuppgifter", ""))
        _set(t[25].cell(i + 2, 3), sg.get("grund", ""))
        _set(t[25].cell(i + 2, 4), sg.get("kommentar", ""))

    _set(t[26].cell(1, 0), s3.get("uppgiftsminimering", ""))

    # T27: necessity — R2 C0 = personuppgift, C1 = nödvändig för att
    _set(t[27].cell(2, 0), s3.get("nodvandighet_personuppgift", ""))
    _set(t[27].cell(2, 1), s3.get("nodvandighet_for", ""))

    _set(t[28].cell(1, 0), s3.get("nodvandighet_text", ""))
    _set(t[29].cell(1, 0), s3.get("lagringsminimering", ""))
    _set(t[30].cell(1, 0), s3.get("sakerhet", ""))
    _set(t[31].cell(1, 0), s3.get("ansvarsskyldighet", ""))
    _set(t[32].cell(1, 0), s3.get("rutiner", ""))
    _set(t[33].cell(1, 0), s3.get("tredjeland", ""))
    _set(t[34].cell(1, 0), s3.get("samlad_bedomning", ""))

    # ── STEG 4 + 5: Riskanalys och Riskhantering (T35) ───────────────────────
    # T35 gets the narrative summaries from both steg4 and steg5.
    # The structured risk data lives in the Excel file (MCP Risks.xlsx).
    s4 = saved.get("steg4", {})
    s4_text = s4.get("sammanfattning", "") if isinstance(s4, dict) else str(s4 or "")
    s5_text = str(saved.get("steg5", "") or "")

    t35_text = s4_text
    if s5_text:
        separator = "\n\nRiskhantering – åtgärder och kvarstående risk:\n"
        t35_text = (t35_text + separator + s5_text) if t35_text else s5_text

    _set(t[35].cell(1, 0), t35_text)

    # Save
    title = saved.get("title", "dpia")
    safe = title.lower().replace(" ", "_").replace("/", "_")
    out_dir = os.path.join(os.path.dirname(__file__), "output")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{safe}_filled.docx")
    doc.save(out_path)
    return out_path
